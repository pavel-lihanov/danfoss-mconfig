import price
import functools
import collections
import operator
import openpyxl
import codecs
import types
import delivery
import wizard

import docx
from docx.oxml.shared import qn

from django.utils.translation import ugettext as _

'''
class PriceNotSpecified(Exception):
    pass

class NotInPricelist(Exception):
    pass
'''

def get_bookmark_par_element(document, bookmark_name):
    """
    Return the named bookmark element
    """
    #doc_element = document._document_part._element
    doc_element = document._element
    bookmarks_list = doc_element.findall('.//' + qn('w:bookmarkStart'))
    for bookmark in bookmarks_list:
        name = bookmark.get(qn('w:name'))
        if name == bookmark_name:
            print (bookmark, type(bookmark))
            #test
            return bookmark
            par = bookmark.getparent()
            if not isinstance(par, docx.oxml.CT_P): 
                raise ValueError('Unexpected parent type', par, type(par))
            else:
                return par
    raise KeyError('Bookmark not found', bookmark_name)
    
def insert_text(par, text):            
    tmp_doc = docx.Document()
    tmp_doc.add_paragraph(text,style='IntenseQuote')
    font = tmp_doc.styles['Normal'].font
    font.name='Calibri'
    bookmark_par_parent = par.getparent()
    index = bookmark_par_parent.index(par) + 1
    for child in tmp_doc._element.body:
        bookmark_par_parent.insert(index, child)
        index = index + 1
    #bookmark_par_parent.remove(par)

class Device:
    '''
    generic device
    subclass to customize price calculation, typecode generation and so on
    '''
    def __init__(self, name, attributes, options, package=False):
        self.name = name
        self.options = dict(options)
        self.attributes = dict(attributes)
        self.is_package = package
        self.template = "mconfig/package.html"
        self.price = None        
        
    def package(self, options, decider):
        new_options = {on: ov for on, ov in options.items()}
        for on, ov in self.options.items():
            if on not in options:               
                #option not specified - select default
                #default selection logic may be altered by user!                
                new_options[on] = decider.select_option(self, on)
                
        return Device(self.name, self.attributes, new_options, package=True)
                    
    def price(self):
        if not self.is_package:     
            raise ValueError('Can only calculate price for packages')
        pr = price.prices[self.name]
        print(self.options)
        pr += sum([price.option_prices[on][ov] for on, ov in self.options.items()])
        return pr
        
    def __repr__(self):
        return 'Device: {0}\n\tattrs{1}\n\toptions{2}\n'.format(self.name, self.attributes, self.options)
        
    def typecode(self):
        return ''
        
VOLTAGES = {
                6000: 'A',
                6600: 'B',
                10000: 'C',
                11000: 'D'
}        

class NotSpecified(Exception):
    pass
        
class VEDADrive(Device):
    class UnknownParameters:
        def __getitem__(self, key):
            raise NotSpecified
    
    class PowerModule:
        def __init__(self, cell, cell_count):
            self.cell = cell
            self.cell_count = cell_count

    class PowerCell:
        def __init__(self, diode_count, nom_voltage, nom_current):
            self.diode_count = diode_count
            self.nom_voltage = nom_voltage
            self.nom_current = nom_current
            
    class SizeDependent:
        def __init__(self, weight, therm_loss):            
            self.weight = weight
            self.therm_loss = therm_loss

    class Frame:
        size_dependent = {}
        def __init__(self, voltage, width, height, length, name='Unknown frame'):
            self.voltage = voltage            
            self.width = width
            self.height = height
            self.length = length                        
            self.size_name = name
            
        @property
        def name(self):
            if self.voltage in VOLTAGES:                                   
                return '{0}{1}{2}'.format(self._type, VOLTAGES[self.voltage], self.size_name)
            else:
                return '{0}{1}'.format(self._type, self.size_name)
        
        def __repr__(self):
            return self.name
            
        def weight(self, device):
            return self.size_dependent[device.attributes['kVA']].weight
        
        def therm_loss(self, device):
            return self.size_dependent[device.attributes['kVA']].therm_loss

        
    class C(Frame):
        _type = 'C'
        service = 'Front and back'
        cooling = 'Air'
        powers = {6000: (315, 400, 500, 630), 10000: (500, 630, 800, 1000)}               
        
        def __init__(self, voltage, width, height, length, powers=(), name='Unknown'):
            super().__init__(voltage, width, height, length, name)
            self.powers = self.powers[voltage]
        
        @classmethod
        def matches(cls, package):            
            if package.options['cooling'] == cls.cooling \
                    and package.options['Service access'] == cls.service \
                    and package.options['brake_mode'] != 'Recuperation':
                for f in cls.frames:
                    if f.size_matches(package):
                        return f
            else:
                return None
                    
        def size_matches(self, package):
            return self.voltage == package.options['mains_voltage'] \
                    and package.attributes['kVA'] in self.powers
            
    CA01 = C(6000, 2150, 2400, 1400, name='01')
    CA01.size_dependent = { 315: SizeDependent(3836, 13), 
                            400: SizeDependent(4036, 16),
                            500: SizeDependent(4236, 20),
                            630: SizeDependent(4436, 25),
    }    
    
    CC01 = C(10000, 2150, 2400, 1400, name='01')
    CC01.size_dependent = { 500: SizeDependent(3904, 13), 
                            630: SizeDependent(4104, 16),
                            800: SizeDependent(4304, 20),
                            1000: SizeDependent(4504, 25),
    }    

    C.frames = (CA01, CC01)
        
    class D(Frame):
        service = 'Front and back'
        cooling = 'Air'        
        _type = 'D'
        
        def __init__(self, voltage, width, height, length, powers=(), name='Unknown D'):
            super().__init__(voltage, width, height, length, name)        
            self.powers = powers            

        @classmethod
        def matches(cls, package):            
            if package.options['cooling'] == cls.cooling \
                    and package.options['Service access'] == cls.service\
                    and package.options['brake_mode'] != 'Recuperation':
                for f in cls.frames:
                    if f.size_matches(package):
                        return f
            else:
                return None
                    
        def size_matches(self, package):            
            return self.voltage == package.options['mains_voltage'] \
                    and package.attributes['kVA'] in self.powers
                                
    DA01 = D(6000, 2150, 2400, 1400, powers = (315, 400, 500, 630), name='01')    
    DA01.size_dependent = { 315: SizeDependent(3836, 13), 
                            400: SizeDependent(4036, 16),
                            500: SizeDependent(4236, 20),
                            630: SizeDependent(4436, 25),
    }    
    
    DA02 = D(6000, 3450, 2200, 1600, powers = (800, 1000, 1250), name='02')
    DA02.size_dependent = { 800: SizeDependent(4430, 32), 
                            1000: SizeDependent(4770, 40),
                            1250: SizeDependent(5180, 50),
    }    
    
    DA03 = D(6000, 4150, 2200, 1600, powers = (1600, 1800, 2000, 2250, 2500), name='03')
    DA01.size_dependent = { 1600: SizeDependent(5085, 63), 
                            1800: SizeDependent(5320, 70),
                            2000: SizeDependent(5560, 80),
                            2250: SizeDependent(5830, 90),
                            2500: SizeDependent(6115, 100)
    }    
    
    DA04 = D(6000, 5850, 2400, 1400, powers = (2800, 3200, 3500, 4000), name='04')
    DA04.size_dependent = { 2800: SizeDependent(8840, 113), 
                            3200: SizeDependent(9190, 125),
                            3500: SizeDependent(9690, 140),
                            4000: SizeDependent(10290, 160),
    }    

    DA05 = D(6000, 7350, 2400, 1400, powers = (4500, 5000), name='05')
    DA05.size_dependent = { 4500: SizeDependent(12700, 180), 
                            5000: SizeDependent(13200, 200),
    }    
    
    DA06 = D(6000, 7650, 2400, 1600, powers = (6300, ), name='06')
    DA06.size_dependent = { 6300: SizeDependent(14000, 250), 
    }    
    
    DA07 = D(6000, 11250, 2400, 1600, powers = (7000, 7900, 8250), name='07')
    DA07.size_dependent = { 7000: SizeDependent(24590, 280), 
                            7900: SizeDependent(26180, 315),
                            8250: SizeDependent(26780, 330),
    }    
    
    DB01 = D(6600, 2150, 2400, 1400, powers = (315, 400, 500, 630), name='01')
    DB01.size_dependent = { 315: SizeDependent(3836, 19), 
                            400: SizeDependent(4036, 19),
                            500: SizeDependent(4236, 22),
                            630: SizeDependent(4436, 28),
    }    
    
    DB02 = D(6600, 3450, 2202, 1600, powers = (800, 1000, 1250), name='02')
    DB02.size_dependent = { 800: SizeDependent(4430, 35), 
                            1000: SizeDependent(4770, 44),
                            1250: SizeDependent(5180, 60),
    }    
    
    DB03 = D(6600, 4150, 2202, 1600, powers = (1600, 1800, 2000, 2250, 2500), name='03')
    DB03.size_dependent = { 1600: SizeDependent(5085, 70), 
                            1800: SizeDependent(5320, 80),
                            2000: SizeDependent(5560, 90),
                            2250: SizeDependent(5830, 100),
                            2500: SizeDependent(6115, 113)
    }    
        
    DB04 = D(6600, 5850, 2400, 1400, powers = (2800, 3200, 3500, 4000), name='04')
    DB04.size_dependent = { 2800: SizeDependent(8840, 125), 
                            3200: SizeDependent(9190, 140),
                            3500: SizeDependent(9690, 155),
                            4000: SizeDependent(10290, 180),
    }    
    
    DB05 = D(6600, 7350, 2400, 1400, powers = (4500, ), name='05')
    DB05.size_dependent = { 4500: SizeDependent(12486, 195), 
    }    
    
    DB06 = D(6600, 7650, 2400, 1600, powers = (5000, 6300, 7000), name='06')
    DB06.size_dependent = { 5000: SizeDependent(12986, 228), 
                            6300: SizeDependent(14348, 250),
                            7000: SizeDependent(14348, 275),
    }    
    
    DB07 = D(6600, 11250, 2400, 1600, powers = (7900, 8250, 9000), name='07')
    DB07.size_dependent = { 7900: SizeDependent(24580, 300), 
                            8250: SizeDependent(26180, 345),
                            9000: SizeDependent(26780, 370),
    }    
    
    
    DC01 = D(10000, 4000, 2000, 1400, powers = (500, 630, 800, 1000), name='01')
    DC01.size_dependent = { 500: SizeDependent(3800, 20), 
                            630: SizeDependent(4000, 25),
                            800: SizeDependent(4250, 32),
                            1000: SizeDependent(4500, 40),
    }    
    
    DC02 = D(10000, 4300, 2200, 1600, powers = (1250, 1600, 1800, 2000, 2250), name='02')
    DC02.size_dependent = { 1250: SizeDependent(5470, 50), 
                            1600: SizeDependent(5910, 63),
                            1800: SizeDependent(6210, 70),
                            2000: SizeDependent(6380, 80),
                            2250: SizeDependent(6670, 90),
    }    

    DC03 = D(10000, 4750, 2250, 1600, powers = (2500, 2800, 3200, 3500, 4000), name='03')
    DC03.size_dependent = { 2500: SizeDependent(6610, 100), 
                            2800: SizeDependent(7105, 113),
                            3200: SizeDependent(7545, 125),
                            3500: SizeDependent(7860, 140),
                            4000: SizeDependent(8375, 160),
    }    
    
    DC04 = D(10000, 7400, 2400, 1600, powers = (4500, 5000, 5500, 6300), name='04')
    DC04.size_dependent = { 4500: SizeDependent(12440, 180), 
                            5000: SizeDependent(13040, 200),
                            5500: SizeDependent(13740, 225),
                            6300: SizeDependent(14340, 250),
    }    
    
    DC05 = D(10000, 8700, 2600, 1600, powers = (7000, ), name='05')
    DC05.size_dependent = { 7000: SizeDependent(16926, 275), 
    }    
    
    DC06 = D(10000, 13300, 2400, 1600, powers = (7900, 8250), name='06')
    DC06.size_dependent = { 7900: SizeDependent(29490, 315), 
                            8250: SizeDependent(32090, 355),
    }    
    
    DC07 = D(10000, 13900, 2400, 1600, powers = (10000, ), name='07')
    DC07.size_dependent = { 10000: SizeDependent(35490, 400), 
    }    
    
    DC08 = D(10000, 14550, 2600, 1600, powers = (12500, ), name='08')
    DC08.size_dependent = { 12500: SizeDependent(42052, 500), 
    }    
    

    DD01 = D(11000, 4000, 2000, 1400, powers = (500, 630, 800, 1000), name='01')
    DD01.size_dependent = { 500: SizeDependent(3800, 24), 
                            630: SizeDependent(4000, 31),
                            800: SizeDependent(4250, 37),
                            1000: SizeDependent(4500, 47),
    }    
    
    DD02 = D(11000, 4300, 2200, 1600, powers = (1600, 1800, 2000, 2250, 2500), name='02')
    DD02.size_dependent = { 1600: SizeDependent(5470, 60), 
                            1800: SizeDependent(5910, 73),
                            2000: SizeDependent(6210, 78),
                            2250: SizeDependent(6380, 88),
                            2500: SizeDependent(6670, 100),
    }    
    
    DD03 = D(11000, 4750, 2250, 1600, powers = (2800, 3200, 3500, 4000), name='03')
    DD03.size_dependent = { 2800: SizeDependent(6610, 120), 
                            3200: SizeDependent(7150, 125),
                            3500: SizeDependent(7545, 148),
                            4000: SizeDependent(7860, 158),
    }    
    
    DD04 = D(11000, 7400, 2400, 1600, powers = (4500, 5000, 5500, 6300, 7000), name='04')
    DD04.size_dependent = { 4500: SizeDependent(8375, 185), 
                            5000: SizeDependent(12590, 200),
                            5500: SizeDependent(13140, 233),
                            6300: SizeDependent(13840, 250),
                            7000: SizeDependent(14490, 278),
    }    
    
    DD05 = D(11000, 8700, 2600, 1600, powers = (7900, ), name='05')
    DD05.size_dependent = { 7900: SizeDependent(15090, 305), 
    }    
    
    DD06 = D(11000, 13300, 2400, 1600, powers = (8250, ), name='06')
    DD06.size_dependent = { 8250: SizeDependent(29490, 353), 
    }    
    
    DD07 = D(11000, 13900, 2400, 1600, powers = (10000, 12500), name='07')
    DD07.size_dependent = { 10000: SizeDependent(32090, 380), 
                            12500: SizeDependent(32090, 460),
    }    
    
    DD08 = D(11000, 14550, 2600, 1600, powers = (14500, ), name='08')
    
    D.frames = (DA01, DA02, DA03, DA04, DA05, DA06, DA07,
                DB01, DB02, DB03, DB04, DB05, DB06, DB07,
                DC01, DC02, DC03, DC04, DC05, DC06, DC07, DC08,
                DD01, DD02, DD03, DD04, DD05, DD06, DD07, DD08,
    )
        
    class R(Frame):
        _type = 'R'
        brake_mode = 'Recuperation'     
        cooling = 'Air'
        service = 'Front and back'
        
        def __init__(self, voltage, width, height, length, motor_type='Induction', powers=(), name='Unknown R'):
            super().__init__(voltage, width, height, length, name)        
            self.powers = powers        
            self.motor_type = motor_type
        
        @classmethod
        def matches(cls, package):
            if package.options['cooling'] == cls.cooling \
                    and package.options['Service access'] == cls.service \
                    and package.options['brake_mode'] == cls.brake_mode:
                for f in cls.frames:
                    if f.size_matches(package):
                        return f

        def size_matches(self, package):
            return self.voltage == package.options['mains_voltage'] \
                    and package.options['motor_type'] == self.motor_type \
                    and package.attributes['kVA'] in self.powers
                                                
    RA01 = R(6000, 2100, 2300, 1400, motor_type='Induction', powers=(400, 500), name='01')
    RA01.size_dependent = { 400: SizeDependent(3150, 23), 
                            500: SizeDependent(3400, 28),
    }    
    
    RA02 = R(6000, 5100, 2100, 1200, motor_type='Induction', powers=(630, 800, 1000), name='02')
    RA02.size_dependent = { 630: SizeDependent(5625, 35), 
                            800: SizeDependent(6185, 45),
                            1000: SizeDependent(6700, 56),
    }    

    RA03 = R(6000, 5900, 2300, 1400, motor_type='Induction', powers=(1250, 1600, 1800, 2000, 2500), name='03')
    RA03.size_dependent = { 1250: SizeDependent(7470, 70), 
                            1600: SizeDependent(8060, 88),
                            1800: SizeDependent(8570, 98),
                            2000: SizeDependent(9270, 112),
                            2500: SizeDependent(9640, 140),
    }    
    
    RA04 = R(6000, 7825, 2400, 1610, motor_type='Induction', powers=(3200, ), name='04')
    RA04.size_dependent = { 3200: SizeDependent(12980, 175), 
    }    
    
    
    RC01 = R(10000, 2400, 2400, 1600, motor_type='Induction', powers=(500, 630), name='01')
    RC01.size_dependent = { 500: SizeDependent(4060, 28), 
                            630: SizeDependent(4260, 35),
    }    
    
    RC02 = R(10000, 5850, 2100, 1200, motor_type='Induction', powers=(800, 1000, 1250, 1600), name='02')
    RC02.size_dependent = { 800: SizeDependent(8085, 45), 
                            1000: SizeDependent(8260, 56),
                            1250: SizeDependent(8850, 70),
                            1600: SizeDependent(9635, 112),
    }    
    
    RC03 = R(10000, 7400, 2300, 1400, motor_type='Induction', powers=(2250, 2500, 3200, 4000), name='03')
    RC03.size_dependent = { 2250: SizeDependent(10855, 126), 
                            2500: SizeDependent(12115, 140),
                            3200: SizeDependent(13205, 175),
                            4000: SizeDependent(14035, 224),
    }    
    
    RAS1 = R(6000, 7275, 2400, 1600, motor_type='PM', powers=(1800, 2000), name='S1')
    RAS1.size_dependent = { 1800: SizeDependent(11030, 98), 
                            2000: SizeDependent(11580, 112),
    }    
    
    RAS2 = R(6000, 7825, 2400, 1400, motor_type='PM', powers=(2250, 2500, 3200, 4000), name='S2')
    RAS2.size_dependent = { 2250: SizeDependent(12980, 126), 
                            2500: SizeDependent(13460, 140),
                            3200: SizeDependent(14140, 175),
                            4000: SizeDependent(15230, 224),
    }    
    
    RAS3 = R(6000, 8850, 2400, 1600, motor_type='PM', powers=(4500, ), name='S3')
    RAS3.size_dependent = { 4500: SizeDependent(21770, 252), 
    }    
    
    R.frames = (RA01, RA02, RA03, RA04, 
                RC01, RC02, RC03,
                RAS1, RAS2, RAS3)
    
    class S(Frame):
        _type = 'S'
        service = 'Front'
        cooling = 'Air'

        def __init__(self, voltage, width, height, length, powers=(), name='Unknown S'):
            super().__init__(voltage, width, height, length, name)        
            self.powers = powers
        
        @classmethod
        def matches(cls, package):
            if package.options['cooling'] == cls.cooling \
                    and package.options['Service access'] == cls.service:
                for f in cls.frames:
                    if f.size_matches(package):
                        return f
                    
        def size_matches(self, package):
            return self.voltage == package.options['mains_voltage'] \
                    and package.attributes['kVA'] in self.powers
                    
    SA01 = S(6000, 3000, 1900, 1200, powers=(315, 400, 500, 630), name='01')
    SA01.size_dependent = { 315: SizeDependent(2516, 13), 
                            400: SizeDependent(2676, 16),
                            500: SizeDependent(2851, 20),
                            630: SizeDependent(3186, 25),
    }    
    
    SA02 = S(6000, 3850, 2100, 1200, powers=(800, 1000, 1250), name='02')
    SA02.size_dependent = { 800: SizeDependent(3876, 32), 
                            1000: SizeDependent(4216, 40),
                            1250: SizeDependent(4656, 50),
    }    
    
    SA03 = S(6000, 4500, 2100, 1200, powers=(1600, 1800, 2000, 2250, 2500), name='03')
    SA03.size_dependent = { 1600: SizeDependent(5070, 63), 
                            1800: SizeDependent(5600, 70),
                            2000: SizeDependent(5600, 80),
                            2250: SizeDependent(6170, 90),
                            2500: SizeDependent(6170, 100)
    }    
    

    SC01 = S(10000, 4550, 1900, 1200, powers=(500, 630, 800, 1000), name='01')
    SC01.size_dependent = { 500: SizeDependent(3630, 20), 
                            630: SizeDependent(3860, 25),
                            800: SizeDependent(4170, 32),
                            1000: SizeDependent(4530, 40),
    }    
    
    SC02 = S(10000, 5400, 2100, 1200, powers=(1250, 1600, 2000, 2250), name='02')
    SC02.size_dependent = { 1250: SizeDependent(3876, 63), 
                            1600: SizeDependent(4216, 70),
                            2000: SizeDependent(4656, 80),
                            2250: SizeDependent(4656, 90),
    }    
    
    SC03 = S(10000, 6500, 2300, 1200, powers=(2500, 3200, 4000), name='03')
    SC03.size_dependent = { 2500: SizeDependent(8410, 100), 
                            3200: SizeDependent(9170, 125),
                            4000: SizeDependent(10720, 160),
    }    

    S.frames = (SA01, SA02, SA03,
                SC01, SC02, SC03,
    )
        
    class L(Frame):
        _type = 'L'
        cooling = 'Liquid'      
        def __init__(self, voltage, width, height, length, powers=(), name='Unknown L'):
            super().__init__(voltage, width, height, length, name)        
            self.powers = powers
        
        @classmethod
        def matches(cls, package):
            if package.options['cooling'] == cls.cooling:
                for f in cls.frames:
                    if f.size_matches(package):
                        return f
                
        def size_matches(self, package):            
            return self.voltage == package.options['mains_voltage'] \
                    and package.attributes['kVA'] in self.powers
                    
    LA01 = L(6000, 8000, 2400, 1400, powers=(2800, 3200, 3500, 4000), name='01')    
    LA02 = L(6000, 9400, 2400, 1600, powers=(4500, 5000, 6300), name='02')
    LA03 = L(6000, 10000, 2400, 1600, powers=(7000, 7900, 8250), name='03')
    LA04 = L(6000, 0, 0, 0, powers=(10000, 12500, 14500), name='04')

    LC01 = L(10000, 10400, 2400, 1400, powers=(4000, 4500, 5000, 6300, 7000), name='01')
    LC02 = L(10000, 11250, 2600, 1600, powers=(7900, 8250), name='02')
    LC03 = L(10000, 0, 0, 0, powers=(10000, ), name='03')
    LC04 = L(10000, 0, 0, 0, powers=(12500, ), name='04')
    
    L.frames = (LA01, LA02, LA03, LA04,
                LC01, LC02, LC03, LC04,
    )    
        
    frames = [C, D, R, S, L]
        
    class Addon(Frame):
        pass
        
    class ManualBypass(Addon):
        _type = 'MB'

        def __init__(self, max_current, width, length, name):
            super().__init__(0 , width, 0, length, name)
            self.max_current = max_current

        def weight(self, device):
            return 0
        
        def therm_loss(self, device):
            return 0
            
        @classmethod
        def matches(cls, package):            
            if package.options['power_option'] == 'Manual bypass':
            #select length
                if package.attributes['nom_current'] <= 500:
                    if package.main_cabinet.length <= 1400:
                        return VEDADrive.MB01
                    else:
                        return VEDADrive.MB02
                else:
                    if package.main_cabinet.length <= 1400:
                        return VEDADrive.MB03
                    else:
                        return VEDADrive.MB04

    MB01 = ManualBypass(500, 800, 1400, name='01')
    MB02 = ManualBypass(500, 800, 1600, name='02')
    MB03 = ManualBypass(10000, 1000, 1400, name='03')
    MB04 = ManualBypass(10000, 1000, 1600, name='04')
    
    class Autobypass(Addon):
        _type = 'AB'
        
        def __init__(self, width, length, name):
            super().__init__(0 , width, 0, length, name)            
            
        def weight(self, device):
            return 0
        
        def therm_loss(self, device):
            return 0            
        
        @classmethod
        def matches(cls, package):            
            if package.options['power_option'] == 'Autobypass':
                if package.main_cabinet.length <= 1400:
                    frame = VEDADrive.AB01
                else:
                    frame = VEDADrive.AB02
                fr = cls(frame.width, frame.length, frame.size_name)
                fr.height = package.main_cabinet.height                
                return fr
            else:
                return None
                
    AB01 = Autobypass(1000, 1400, name='01')
    AB02 = Autobypass(1000, 1600, name='02')
    
    class Reactor(Addon):
        _type = 'E'
        
        def __init__(self, voltage, width, length, height, powers, name):
            super().__init__(voltage , width, height, length, name)            
            self.powers = powers
            
        @classmethod
        def matches(cls, package):
            if package.options['Output reactor'] == 'Yes':
                for f in cls.frames:
                    if f.size_matches(package):
                        return f
                
        def size_matches(self, package):            
            return self.voltage == package.attributes['voltage'] \
                    and package.attributes['kVA'] in self.powers
                    
    EA01 = Reactor(6000, 1200, 2400, 1400, powers=(250, 315, 400, 500), name = '01')
    EA01.size_dependent = { 250: SizeDependent(209, 0), 
                            315: SizeDependent(223, 0),
                            400: SizeDependent(256, 0),
                            500: SizeDependent(286, 0),
    }    
    
    EA02 = Reactor(6000, 1400, 2200, 1600, powers=(630, 800, 1000), name = '02')
    EA02.size_dependent = { 630: SizeDependent(335, 0), 
                            800: SizeDependent(360, 0),
                            1000: SizeDependent(480, 0),
    }    
    
    EA03 = Reactor(6000, 1600, 2200, 1600, powers=(1250, 1400, 1600, 1800, 2000), name = '03')
    EA03.size_dependent = { 1250: SizeDependent(566, 0), 
                            1400: SizeDependent(682, 0),
                            1600: SizeDependent(682, 0),
                            1800: SizeDependent(837, 0),
                            2000: SizeDependent(837, 0)
    }    
    
    EA04 = Reactor(6000, 0, 0, 0, powers=(2250, 2500, 2800, 3200), name = '04')

    EC01 = Reactor(10000, 1200, 2400, 1400, powers=(250, 315, 400, 500), name = '01')
    EC01.size_dependent = { 250: SizeDependent(283, 0), 
                            315: SizeDependent(293, 0),
                            400: SizeDependent(339, 0),
                            500: SizeDependent(384, 0),
    }    
    
    EC02 = Reactor(10000, 1400, 2200, 1600, powers=(630, 800, 1000), name = '02')
    EC02.size_dependent = { 630: SizeDependent(0, 0), 
                            800: SizeDependent(0, 0),
                            1000: SizeDependent(0, 0),
    }    
    
    EC03 = Reactor(10000, 1600, 2200, 1600, powers=(1250, 1400, 1600, 1800, 2000), name = '03')
    EC03.size_dependent = { 1250: SizeDependent(759, 0), 
                            1400: SizeDependent(990, 0),
                            1600: SizeDependent(990, 0),
                            1800: SizeDependent(1217, 0),
                            2000: SizeDependent(1217, 0)
    }    
    
    EC04 = Reactor(10000, 0, 0, 0, powers=(2250, 2500, 2800, 3200), name = '04')
    EC04.size_dependent = { 2250: SizeDependent(1408, 0), 
                            2500: SizeDependent(1408, 0),
                            2800: SizeDependent(1691, 0),
                            3200: SizeDependent(1691, 0),
    }    
        
    Reactor.frames = (EA01, EA02, EA03, EA04)            
        
    class Multistart(Addon):
        _type = 'MS'
        def __init__(self, width, length, name):
            super().__init__(0 , width, 0, length, name)            

        @classmethod
        def matches(cls, package):
            print('MS', package.main_cabinet)
            if package.options['power_option'] == 'Multistart':
                motor_count = package.options['multi_motors']
                if motor_count == 2:
                    frame = VEDADrive.MS01
                elif motor_count == 3:
                    frame = VEDADrive.MS02
                elif motor_count == 4:
                    frame = VEDADrive.MS03
                if frame:                    
                    fr = cls(frame.width, frame.length, frame.size_name)
                    fr.height = package.main_cabinet.height
                    return fr
                else:
                    return None
            else:
                return None
                
    MS01 = Multistart(800, 1600, '01')
    MS02 = Multistart(800, 1600, '02')
    MS03 = Multistart(800, 1600, '03')
                
    Multistart.frames = (MS01, MS02, MS03)
        
    addons = [ManualBypass, Autobypass, Reactor, Multistart]
    
        
    def __init__(self, name, attributes, options=None, package=False):
        Device.__init__(self, name, {}, {}, package)
        self._attrs = attributes
        self.attributes = VEDAAttrs(self)
        self.template = "mconfig/VEDADrive.html"
        
        if options:
            self.options = VEDAOpts(device=self, **options)
        else:
            self.options = VEDAOpts(device=self)

        #bookmarks getters
        self.bookmarks = {   'device_price': self.get_price,
                        'type_code': self.order_code,
                        #'corpus': self.corpus
                        
                        }

    #getters for offer
    def get_price(self):
        assert(self.is_package)
        return self.price.total
    
    def corpus(self):
        assert(self.is_package)
        corpus=res1['Корпус']  
        return corpus

    def input_cable(self):
        assert(self.is_package)
        input_cable=res1['Ввод питающего кабеля'] 
        return input_cable
        
    def output_cable(self):
        assert(self.is_package)
        output_cable=res1['Вывод кабеля двигателя'] 
        return output_cable
        
    def service(self):
        assert(self.is_package)
        output_cable=res1['Обслуживание'] 
        return service
    
    def opt_b(self):
        assert(self.is_package)
        opt_b=res1['Опция B'] 
        return opt_b
        
    #выходное_напряжение  
    def output_voltage(self):
        assert(self.is_package)
        output_voltage=self.attributes['voltage'] 
        return output_voltage
    #входное_напряжение     
    def input_voltage(self):
        assert(self.is_package)
        input_voltage=wizard.opts['mains_voltage'] 
        return input_voltage    
        
    def power_cell(self):
        assert(self.is_package)
        power_cell=res1['Силовых ячеек на фазу'] 
        return power_cell
        
    def nom_current(self):
        assert(self.is_package)
        nom_current=self.attributes['nom_current'] 
        return nom_current
    #полная мощность
    def kVA(self):
        assert(self.is_package)
        kVA=self.attributes['kVA'] 
        return kVA
    
    def order_code(self):
        assert(self.is_package)
        fields = [(0, self.name), (17, '{0:03}'.format(self.attributes['nom_current']))] \
            + [(self.options._opts[n].code_pos, self.options._opts[n].choices_to_codes[v]) for n,v in self.options.items() if n in self.options._opts]
        return ''.join([f[1] for f in sorted(fields, key=operator.itemgetter(0))])
        
    #cabinet calculation logic
    def get_power_module(self):
        return None
    
    def get_frame(self):
        assert(self.is_package)
        #service = self.options['Service access']
        #current = self.attributes['nom_current']
        #cooling = self.options['cooling']
        #recuperation = self.options['brake_mode']
        #voltage = self.attributes['voltage']        
        for f in VEDADrive.frames:
            frame = f.matches(self)
            if frame:
                print('Frame: {0}'.format(frame))
                return frame
        raise ValueError('No matching frame')
        
    def get_addons(self):
        res = []
        print(self.main_cabinet)
        for f in VEDADrive.addons:
            a = f.matches(self)
            if a:
                res.append(a)
        return res
        
    #make package from options
    def package(self, options, decider):
        #check if we need power transformer                    
        options = dict(options)
        motor_voltage = self.attributes['voltage']
        if 'mains_voltage' in options:
            mains_voltage = options['mains_voltage']
            print('Mains voltage:', mains_voltage)
            print('Motor voltage:', motor_voltage)
            if motor_voltage != mains_voltage:                
                options['transformer'] =  str(mains_voltage) + '-' + str(motor_voltage)                
        else:
            raise KeyError('Supply voltage is not specified')
        
        new_options = collections.OrderedDict()
        
        for on in self.options._opts:                
            if on not in options:
                #option not specified - select default
                #default selection logic may be altered by user!                
                new_options[on] = decider.select_option(self, on)
            else:
                new_options[on] = options[on]
                
        for n,v in options.items():
            if n not in new_options:
                new_options[n] = v
              
        #calculate frames
        #main_cabinet = self.get_frame(new_options)
        #addons = self.get_addons(new_options)
        pkg = VEDADrive(self.name, self.attributes, new_options, package=True)
        
        pkg.power_module = pkg.get_power_module()
        pkg.main_cabinet = pkg.get_frame()
        pkg.addons = pkg.get_addons()
        
        return pkg
       
    def calculate_price(self):
        assert(self.is_package)
        pricelist = price.price_lists['VEDADrive']
        self.price = pricelist.get_price(self)
        _, self.price.delivery_cost, _ = delivery.get_delivery_details(self)
        #calculate delivery costs        
                
    def make_offer_template(self, path):            
        assert(self.is_package)
        assert(self.price)
        assert(self.corpus)
        assert(self.input_cable)
        assert(self.output_cable)
        assert(self.service)
        assert(self.opt_b)
        assert(self.output_voltage)
        assert(self.input_voltage)
        assert(self.power_cell)
        assert(self.nom_current)
        assert(self.kVA)
        
        doc = docx.Document('offer_template.docx')
        par = get_bookmark_par_element(doc, "order_code")
        #insert_text(par, self.order_code())
        
        #for b,g in self.bookmarks.items():
        #    par = get_bookmark_par_element(doc, b)
        #    insert_text(par, g())
        #par=doc.add_heading('Введение',level=1)
        #insert_text(par,self)
        
        #par = get_bookmark_par_element(doc, "price")
        #insert_text(par, '{0:.2f}'.format(self.price.total))
        
        #par = get_bookmark_par_element(doc, "price1")
        #insert_text(par, '{0:.2f}'.format(self.price.total))
        
        par = get_bookmark_par_element(doc, "order_code")
        insert_text(par, self.order_code())
        
        par = get_bookmark_par_element(doc, "order_code1")
        insert_text(par, self.order_code())
        
        par = get_bookmark_par_element(doc, "order_code2")
        insert_text(par, self.order_code())
        
        par = get_bookmark_par_element(doc, "order_code3")
        insert_text(par, self.order_code())
        
        par = get_bookmark_par_element(doc, "order_code4")
        insert_text(par, self.order_code())
        
        par = get_bookmark_par_element(doc, "order_code5")
        insert_text(par, self.order_code())
        
        #par = get_bookmark_par_element(doc, "corpus")
        #insert_text(par, self.corpus())
        #par = get_bookmark_par_element(doc, "order_code6")
        #insert_text(par, self.order_code())
        
        par = get_bookmark_par_element(doc, "corpus1")
        insert_text(par, self.corpus())
        
        par = get_bookmark_par_element(doc, "input_cable")
        insert_text(par, self.input_cable())
        
        par = get_bookmark_par_element(doc, "output_cable")
        insert_text(par, self.output_cable())
        
        par = get_bookmark_par_element(doc, "service")
        insert_text(par, self.output_cable())
        
        
        par = get_bookmark_par_element(doc, "opt_b")
        if self.opt_b() == 'Нет':
            insert_text(par, 'ModBus RTU')    
        else:
            insert_text(par, self.opt_b())

        par = get_bookmark_par_element(doc, "opt_b1")
        if self.opt_b() == 'Нет':
            insert_text(par, 'ModBus-встроен')    
        else:
            insert_text(par, self.opt_b()) 
            
        par = get_bookmark_par_element(doc, "power_cell")
        insert_text(par,'{0:.0f}'.format(self.power_cell()*3)+' ( ' +'{0:.0f}'.format(self.power_cell())+' на одну фазу)')
        #insert_text(par, ('{0:.0f}'.format(self.power_cell()*3)+'( '+ '{0:.0f}'.format(self.power_cell()) +' на одну фазу)'))        
            
        par = get_bookmark_par_element(doc, "output_voltage")
        insert_text(par, ('{0:.0f}'.format(self.output_voltage()/1000)))        
        
        par = get_bookmark_par_element(doc, "input_voltage")
        insert_text(par, ('{0:.0f}'.format(self.input_voltage()/1000))) 
        
        par = get_bookmark_par_element(doc, "output_voltage_trans")
        insert_text(par,'{0:.0f}'.format(self.power_cell()*3)+' x 690')
        
        par = get_bookmark_par_element(doc, "diods")
        insert_text(par,'{0:.0f}'.format(self.power_cell()*3*2))
        
        par = get_bookmark_par_element(doc, "nom_current")
        insert_text(par,'690В, '+'{0:.0f}'.format(self.nom_current())+' A')
        
        par = get_bookmark_par_element(doc, "height")
        insert_text(par,'{0:.0f}'.format(self.main_cabinet.height))
        
        par = get_bookmark_par_element(doc, "length")
        insert_text(par,'{0:.0f}'.format(self.main_cabinet.length))
        
        par = get_bookmark_par_element(doc, "width")
        insert_text(par,'{0:.0f}'.format(self.main_cabinet.width + sum([o.width for o in self.addons])))
        
        par = get_bookmark_par_element(doc, "weight")
        insert_text(par,'{0:.0f}'.format(self.main_cabinet.weight(self) + sum([o.weight(self) for o in self.addons])))
 
        par = get_bookmark_par_element(doc, "kVA")
        insert_text(par,'{0:.0f}'.format(self.kVA()))
        
        doc.save(path)        
        
    def short_descr(self):
        sacc = self.options._opts['Service access'].display_choices[self.options['Service access']]
        sacc = sacc[0].lower() + sacc[1:]
        return _('VEDADRIVE Uin={0}kV, Uout={5}kV, {1}A {6} cells per phase, {2}, {3} {4} service access').format(
            self.options['mains_voltage']//1000,
            self.attributes['nom_current'],
            _('Air-cooled') if self.options['cooling']=='Air' else _('Liquid_cooled'),
            self.options['enclosure'], 
            sacc,
            #self.options['Service access'], 
            self.attributes['voltage']//1000,
            self.options['power_cells'])

    #do not use, use PriceList.package_matches_pricelist()
    def matches_pricelist(self, row):
        price_name = '{0}-{1}{2}'.format(row[1].value, row[2].value, row[4].value)
        if self.name != price_name:
            return False
            
        for name,val in self.options._opts.items():
            if not val.matches_pricelist(self.options[name], row):
                return False
            
        return True
        
    @property
    def width(self):
        assert self.is_package
        return self.main_cabinet.width + sum([o.width for o in self.addons])
        
    @property
    def height(self):
        assert self.is_package
        return self.main_cabinet.height
    
    @property
    def length(self):
        assert self.is_package
        return self.main_cabinet.length
        
    @property
    def weight(self):
        assert self.is_package
        return self.main_cabinet.weight(self) + sum([o.weight(self) for o in self.addons])

    @property
    def therm_loss(self):
        assert self.is_package
        return self.main_cabinet.therm_loss(self) + sum([o.therm_loss(self) for o in self.addons])        
        
    def delivery_items(self):
        #TODO: get weight (and mb size) of drive components
        #see get_frame() and get_addons() methods
        assert(self.is_package)
        return [15, ]
        
    def display_options(self):
        infos = self.options._opts
        opts = self.options.items()        
        
        rl = []        
        
        il = [k for k,v in infos.items()]
        
        i = len(il)
        
        for k,v in opts:
            if k in infos:
                #res[infos[k].display_name] = infos[k].display_choices[v]
                rl.append((il.index(k), (infos[k].display_name, infos[k].display_choices[v])))            
            else:
                #res[k] = v
                print('Unknown option?', k, v)
                rl.append((i, (k, v)))
                i+=1 
        global res1        
        res1 = collections.OrderedDict([a[1] for a in sorted(rl, key=lambda a: a[0])])
        print ('result=')
        print (res1)
        
        #corpus=res['Корпус']
        #print (corpus)
        
        return res1
        
        
class VEDAAttrs:
    def __init__(self, drive):
        self.drive = drive
        
    def __getitem__(self, key):
        if key == 'overload_current':
            return self.drive._attrs['nom_current'] * 1.2
        else:
            return self.drive._attrs[key]
            
class VEDAOption:
    def __init__(self, name, choices, codes, code_pos, display_name=None, display_choices=None, price_field=None, price_getter=None):
        self.name = name
        self._choices = choices
        self.codes = codes
        self.code_pos = code_pos
        self.choices_to_codes = dict(zip(choices, codes))
        self.codes_to_choices = dict(zip(codes, choices))
        self.price_field = price_field
        self.price_getter = price_getter        
        self.display_name = display_name if display_name else name
        self.display_choices = display_choices if display_choices else {a:a for a in choices}        
            
    def choices(self, device):
        return self._choices
           
    def matches_pricelist(self, option_value, row):        
        '''
        TODO: should be def matches_pricelist(self, option_value, entry):
            return entry.get_option(self.name) == option_value
        '''
        if self.price_getter is not None:            
            price_opt = self.price_getter(self, row[self.price_field].value)
            return price_opt == option_value                
        else:
            return True        
        
    def generic_getter(self, value):
        return self.codes_to_choices[str(value)]
        
    def letter_getter (self, value):
        prefix = self.codes[0][0]
        try:
            value = value.replace('Х', 'X')
            return self.codes_to_choices[prefix + str(value)]
        except KeyError:
            print ('Unknown option code', prefix + str(value))
            print (codecs.encode(prefix + str(value), 'unicode_escape'))
            print ('Have', self.codes_to_choices)
            raise
        
    def mains_getter(self, value):
        return self.codes_to_choices[str(value)[0]]
        
    def motor_getter(self, value):
        return self.codes_to_choices[str(value)[1]]
        
    def poweropt_getter(self, value):
        if value=='XX':
            return self.codes_to_choices['AX']
        else:
            return self.codes_to_choices['A'+str(value)[0]]
            
class CoolingOption(VEDAOption):
    def choices(self, dev):
        if dev.attributes['voltage'] == 6000:
            if dev.attributes['nom_current'] >= 275:
                return self._choices
            else:
                return (self._choices[0], )
        elif dev.attributes['voltage'] == 10000:
            if dev.attributes['nom_current'] >= 260:
                return self._choices
            else:
                return (self._choices[0], )    

class PowerCellOption(VEDAOption):
    def choices(self, dev):
        if dev.attributes['voltage'] == 6000:
            #6kV - 5 or 6 cells
            return self._choices[0:2]
        elif dev.attributes['voltage'] == 10000:
            return self._choices[2:]
                
                
class ServiceOption(VEDAOption):                
    def choices(self, dev):
        if dev.attributes['nom_current'] <= 243:
                return self._choices
        else:
            return (self._choices[1], )    

class VEDAOpts:
    def __init__(self, **args):        
        self._opts = collections.OrderedDict()
        for o in (
                    VEDAOption('motor_type', ('Induction', 'PM'), ('A', 'S'), 14, 
                        price_field=7, price_getter=VEDAOption.generic_getter,
                        display_name = _('Motor type'), display_choices = {'Induction':_('Induction'), 'PM':_('PMSM')},
                        ),
                    VEDAOption('input_freq', (50, 60), ('F5', 'F6'), 10, 
                        price_field=5, price_getter=VEDAOption.generic_getter,
                        display_name = _('Supply frequency'), 
                        ),
                    VEDAOption('enclosure', ('IP30', 'IP31', 'IP41', 'IP42', 'IP54'), ('30', '31', '41', '42', '54'), 12, 
                        price_field=6, price_getter=VEDAOption.generic_getter,
                        display_name = _('Enclosure'), 
                        ),
                    VEDAOption('control_mode', ('U/f', 'Vector control'), ('S', 'V'), 15, 
                        price_field=8, price_getter=VEDAOption.generic_getter,
                        display_name = _('Motor contol mode'), display_choices={'U/f':_('U/f control'), 'Vector control':_('Vector control')}
                        ),
                    VEDAOption('brake_mode', ('Coast', 'Dynamic', 'Recuperation'), ('X', 'B', 'R'), 16, 
                        price_field=9, price_getter=VEDAOption.generic_getter,
                        display_name = _('Braking mode'), display_choices={'Coast':_('Coasting stop'), 'Dynamic':_('Dynamic braking'), 'Recuperation':_('Recuperation')}
                        ),
                    CoolingOption('cooling', ('Air', 'Liquid'), ('A', 'L'), 20, 
                        price_field=11, price_getter=VEDAOption.generic_getter,
                        display_name= _('Cooling'), display_choices={'Air':_('Air'), 'Liquid':_('Liquid')}
                        ),
                    VEDAOption('power_cell_autobypass', ('No', 'Yes'), ('X', 'C'), 21, 
                        price_field=12, price_getter=VEDAOption.generic_getter,
                        display_name=_('Power cell autobypass'), display_choices={'No':_('No'), 'Yes':_('Yes')}),
                    VEDAOption('power_option', ('None', 'Manual bypass', 'Autobypass', 'Multistart', 'Interchange'), ('AX', 'A2', 'A1', 'A3', 'A4'), 22, 
                        price_field=13, price_getter=VEDAOption.poweropt_getter,
                        display_name=_('Power option'), display_choices={'None':_('None'), 'Manual bypass':_('Manual bypass'), 'Autobypass':_('Autobypass'), 'Multistart':_('Multistart'), 'Interchange':_('Interchange')},
                        ),
                    VEDAOption('multi_motors', (1, 2, 3, 4), ('1', '2', '3', '4'), 24, 
                        price_field=13, price_getter=None,
                        display_name=_('Number of motors')),
                    VEDAOption('fieldbus', ('None', 'Encoder', 'EtherNet IP', 'ProfiBus DP', 'Modbus TCP/IP'), ('BX', 'B1', 'B2', 'B3', 'B4'), 25, 
                        price_field=14, price_getter=VEDAOption.letter_getter,
                        display_name=_('Option B'), display_choices={'None':_('None'), 'Encoder':_('Encoder board'), 'EtherNet IP':_('EtherNet IP'), 'ProfiBus DP':_('ProfiBus DP'), 'Modbus TCP/IP':_('Modbus TCP/IP')}
                        ),
                    VEDAOption('transformer', ('None', '10000-6000', '10000-6600', '6000-10000', '6600-10000'), ('CX', 'C1', 'C2', 'C3', 'C4'), 27, 
                        price_field=15, price_getter=VEDAOption.letter_getter,
                        display_name=_('Power transformer'), display_choices={'None':_('None'), '10000-6000':'10000-6000', '10000-6600':'10000-6600', '6000-10000':'6000-10000', '6600-10000':'6600-10000'}
                        ),
                    VEDAOption('PMSM exciter', ('No', 'Yes'), ('DX', 'D1'), 29, 
                        price_field=16, price_getter=VEDAOption.letter_getter,
                        display_name=_('PMSM exciter'), display_choices={'No':_('No'), 'Yes':_('Yes')}
                        ),
                    VEDAOption('Input mains location', ('Top', 'Bottom'), ('2', '1'), 31, 
                        price_field=17, price_getter=VEDAOption.mains_getter,
                        display_name=_('Input mains location'), display_choices={'Top':_('Top'), 'Bottom':_('Bottom')}
                        ),
                    VEDAOption('Motor cable location', ('Top', 'Bottom'), ('2', '1'), 32, 
                        price_field=17, price_getter=VEDAOption.motor_getter,
                        display_name=_('Motor cable location'), display_choices={'Top':_('Top'), 'Bottom':_('Bottom')}
                        ),
                    VEDAOption('Output reactor', ('No', 'Yes'), ('EX', 'E1'), 33, 
                        price_field=18, price_getter=VEDAOption.letter_getter,
                        display_name=_('Output reactor'), display_choices={'No':_('No'), 'Yes':_('Yes')}
                        ),
                    ServiceOption('Service access', ('Front', 'Front and back'), ('S', 'D'), 35, 
                        price_field=19, price_getter=VEDAOption.generic_getter,
                        display_name=_('Service access'), display_choices={'Front':_('Front'), 'Front and back':_('Front and back')}
                        ),                    
                    VEDAOption('mains_voltage', (6000, 10000), ('', ''), 0,
                        display_name=_('Motor voltage'),
                        ),                                                
                    #motor_voltage is attributes['voltage']
                    #VEDAOption('motor_voltage', (6000, 10000), ('', ''), 0,
                    #    display_name=_('Motor voltage'),
                    #    ),                        
                    PowerCellOption('power_cells', (5, 6, 8, 9), ('', '', '', ''), 0,
                        display_name=_('Power cells count'),
                        ),
                        
                ):
            self._opts[o.name] = o
        
        if 'device' in args:
            dev = args['device'] if 'device' in args else None
            del args['device']
            
        self.opts = collections.OrderedDict()
        for on, ov in self._opts.items():
                if dev:                    
                    self.opts[on] = ov.choices(dev)
                else:
                    self.opts[on] = ()
                        
        self.opts.update(args)

    def __getitem__(self, key):
        return self.opts[key]
        
        
    def infos(self):
        return self._opts.items()
        
    def items(self):
        return self.opts.items()

devices = [
    VEDADrive('VD-P315KU1', attributes={'voltage': 6000, 'nom_current': 31, 'kVA': 315}),
    VEDADrive('VD-P400KU1', attributes={'voltage': 6000, 'nom_current': 40, 'kVA': 400}),
    VEDADrive('VD-P500KU1', attributes={'voltage': 6000, 'nom_current': 48, 'kVA': 500}),
    VEDADrive('VD-P630KU1', attributes={'voltage': 6000, 'nom_current': 61, 'kVA': 630}),
    VEDADrive('VD-P800KU1', attributes={'voltage': 6000, 'nom_current': 77, 'kVA': 800}),
    VEDADrive('VD-P1000U1', attributes={'voltage': 6000, 'nom_current': 96, 'kVA': 1000}),
    VEDADrive('VD-P1250U1', attributes={'voltage': 6000, 'nom_current': 130, 'kVA': 1250}),
    VEDADrive('VD-P1600U1', attributes={'voltage': 6000, 'nom_current': 154, 'kVA': 1600}),
    VEDADrive('VD-P1800U1', attributes={'voltage': 6000, 'nom_current': 173, 'kVA': 1800}),
    VEDADrive('VD-P2000U1', attributes={'voltage': 6000, 'nom_current': 192, 'kVA': 2000}),
    VEDADrive('VD-P2250U1', attributes={'voltage': 6000, 'nom_current': 230, 'kVA': 2250}),
    VEDADrive('VD-P2500U1', attributes={'voltage': 6000, 'nom_current': 243, 'kVA': 2500}),
    VEDADrive('VD-P2800U1', attributes={'voltage': 6000, 'nom_current': 275, 'kVA': 2800}),
    VEDADrive('VD-P3200U1', attributes={'voltage': 6000, 'nom_current': 304, 'kVA': 3200}),
    VEDADrive('VD-P3500U1', attributes={'voltage': 6000, 'nom_current': 340, 'kVA': 3500}),
    VEDADrive('VD-P4000U1', attributes={'voltage': 6000, 'nom_current': 400, 'kVA': 4000}),
    VEDADrive('VD-P4500U1', attributes={'voltage': 6000, 'nom_current': 425, 'kVA': 4500}),
    VEDADrive('VD-P5000U1', attributes={'voltage': 6000, 'nom_current': 500, 'kVA': 5000}),
    VEDADrive('VD-P6300U1', attributes={'voltage': 6000, 'nom_current': 600, 'kVA': 6300}),
    VEDADrive('VD-P7000U1', attributes={'voltage': 6000, 'nom_current': 660, 'kVA': 7000}),
    VEDADrive('VD-P7900U1', attributes={'voltage': 6000, 'nom_current': 750, 'kVA': 7900}),
    VEDADrive('VD-P8250U1', attributes={'voltage': 6000, 'nom_current': 800, 'kVA': 8250}),
    
    VEDADrive('VD-P500KU3', attributes={'voltage': 10000, 'nom_current': 31, 'kVA': 500}, options={'cooling': 'Air'}),
    VEDADrive('VD-P630KU3', attributes={'voltage': 10000, 'nom_current': 40, 'kVA': 630}, options={'cooling': 'Air'}),
    VEDADrive('VD-P800KU3', attributes={'voltage': 10000, 'nom_current': 48, 'kVA': 800}, options={'cooling': 'Air'}),
    VEDADrive('VD-P1000U3', attributes={'voltage': 10000, 'nom_current': 61, 'kVA': 1000}, options={'cooling': 'Air'}),
    VEDADrive('VD-P1250U3', attributes={'voltage': 10000, 'nom_current': 77, 'kVA': 1250}, options={'cooling': 'Air'}),
    VEDADrive('VD-P1600U3', attributes={'voltage': 10000, 'nom_current': 96, 'kVA': 1600}, options={'cooling': 'Air'}),
    VEDADrive('VD-P1800U3', attributes={'voltage': 10000, 'nom_current': 104, 'kVA': 1800}, options={'cooling': 'Air'}),
    VEDADrive('VD-P2000U3', attributes={'voltage': 10000, 'nom_current': 115, 'kVA': 2000}, options={'cooling': 'Air'}),
    VEDADrive('VD-P2250U3', attributes={'voltage': 10000, 'nom_current': 130, 'kVA': 2250}, options={'cooling': 'Air'}),
    VEDADrive('VD-P2500U3', attributes={'voltage': 10000, 'nom_current': 154, 'kVA': 2500}, options={'cooling': 'Air'}),
    VEDADrive('VD-P2800U3', attributes={'voltage': 10000, 'nom_current': 165, 'kVA': 2800}, options={'cooling': 'Air'}),
    VEDADrive('VD-P3150U3', attributes={'voltage': 10000, 'nom_current': 192, 'kVA': 3150}, options={'cooling': 'Air'}),
    VEDADrive('VD-P3500U3', attributes={'voltage': 10000, 'nom_current': 205, 'kVA': 3500}, options={'cooling': 'Air'}),
    VEDADrive('VD-P4000U3', attributes={'voltage': 10000, 'nom_current': 243, 'kVA': 4000}, options={'cooling': 'Air'}),
    VEDADrive('VD-P4500U3', attributes={'voltage': 10000, 'nom_current': 260, 'kVA': 4500}, options={'cooling': 'Air'}),
    VEDADrive('VD-P5000U3', attributes={'voltage': 10000, 'nom_current': 304, 'kVA': 5000}),
    VEDADrive('VD-P5630U3', attributes={'voltage': 10000, 'nom_current': 325, 'kVA': 5630}),
    VEDADrive('VD-P6300U3', attributes={'voltage': 10000, 'nom_current': 364, 'kVA': 6300}),
    VEDADrive('VD-P7000U3', attributes={'voltage': 10000, 'nom_current': 400, 'kVA': 7000}),
    VEDADrive('VD-P7900U3', attributes={'voltage': 10000, 'nom_current': 462, 'kVA': 7900}),
    VEDADrive('VD-P8900U3', attributes={'voltage': 10000, 'nom_current': 500, 'kVA': 8900}),
    VEDADrive('VD-P10M0U3', attributes={'voltage': 10000, 'nom_current': 600, 'kVA': 10000}),
    VEDADrive('VD-P12M5U3', attributes={'voltage': 10000, 'nom_current': 800, 'kVA': 12500}),
]
